"""
Document Loader

This module provides document loading functionality for various file formats.
"""

import os
from typing import List, Dict, Optional
from pathlib import Path
import mimetypes


class Document:
    """
    Represents a loaded document.
    """

    def __init__(
        self,
        content: str,
        metadata: Optional[Dict] = None,
        source: Optional[str] = None
    ):
        """
        Initialize a document.

        Args:
            content: Document text content
            metadata: Additional metadata
            source: Source file path
        """
        self.content = content
        self.metadata = metadata or {}
        self.source = source


class DocumentLoader:
    """
    Document loader for various file formats.
    """

    def __init__(self):
        """Initialize document loader."""
        self.supported_extensions = {
            '.txt': self._load_text,
            '.md': self._load_text,
            '.markdown': self._load_text,
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.doc': self._load_doc,
            '.html': self._load_html,
            '.htm': self._load_html,
        }

    def load(self, file_path: str) -> Document:
        """
        Load a single document.

        Args:
            file_path: Path to the file

        Returns:
            Document object

        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file does not exist
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        extension = path.suffix.lower()

        if extension not in self.supported_extensions:
            raise ValueError(
                f"不支持的文件格式: {extension}. "
                f"支持的格式: {', '.join(self.supported_extensions.keys())}"
            )

        # Load content using appropriate loader
        content = self.supported_extensions[extension](path)

        # Create document with metadata
        metadata = {
            'filename': path.name,
            'file_path': str(path.absolute()),
            'file_size': path.stat().st_size,
            'file_type': extension,
        }

        return Document(content=content, metadata=metadata, source=str(path))

    def load_directory(
        self,
        directory_path: str,
        recursive: bool = True,
        file_extensions: Optional[List[str]] = None
    ) -> List[Document]:
        """
        Load all documents from a directory.

        Args:
            directory_path: Path to the directory
            recursive: Whether to search recursively
            file_extensions: List of file extensions to include (e.g., ['.pdf', '.txt'])

        Returns:
            List of Document objects
        """
        directory = Path(directory_path)

        if not directory.exists():
            raise FileNotFoundError(f"目录不存在: {directory_path}")

        if not directory.is_dir():
            raise ValueError(f"不是一个目录: {directory_path}")

        documents = []

        # Determine which extensions to load
        if file_extensions:
            extensions = [ext.lower() for ext in file_extensions]
        else:
            extensions = list(self.supported_extensions.keys())

        # Find all matching files
        pattern = '**/*' if recursive else '*'

        for path in directory.glob(pattern):
            if path.is_file() and path.suffix.lower() in extensions:
                try:
                    doc = self.load(str(path))
                    documents.append(doc)
                except Exception as e:
                    print(f"加载文件失败 {path}: {str(e)}")

        return documents

    def _load_text(self, path: Path) -> str:
        """Load text file."""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(path, 'r', encoding='gbk') as f:
                return f.read()

    def _load_pdf(self, path: Path) -> str:
        """Load PDF file."""
        try:
            import PyPDF2
            
            text = []
            with open(path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text.append(page.extract_text())
            
            return '\n'.join(text)
        except ImportError:
            raise ImportError(
                "需要安装 PyPDF2 来读取 PDF 文件: pip install PyPDF2"
            )

    def _load_docx(self, path: Path) -> str:
        """Load DOCX file."""
        try:
            import docx
            
            doc = docx.Document(path)
            paragraphs = [para.text for para in doc.paragraphs]
            
            return '\n'.join(paragraphs)
        except ImportError:
            raise ImportError(
                "需要安装 python-docx 来读取 DOCX 文件: pip install python-docx"
            )

    def _load_doc(self, path: Path) -> str:
        """Load DOC file (legacy format)."""
        # DOC format is more complex and requires external tools
        raise NotImplementedError(
            "旧版 .doc 格式需要额外的工具支持。请将文件转换为 .docx 格式。"
        )

    def _load_html(self, path: Path) -> str:
        """Load HTML file."""
        try:
            from bs4 import BeautifulSoup
            
            with open(path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                # Extract text from HTML
                return soup.get_text(separator='\n', strip=True)
        except ImportError:
            raise ImportError(
                "需要安装 beautifulsoup4 来读取 HTML 文件: pip install beautifulsoup4"
            )


class TextSplitter:
    """
    Text splitter for chunking documents.
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separator: str = "\n"
    ):
        """
        Initialize text splitter.

        Args:
            chunk_size: Maximum size of each chunk in characters
            chunk_overlap: Number of characters to overlap between chunks
            separator: Text separator for splitting
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separator = separator

    def split_text(self, text: str) -> List[str]:
        """
        Split text into chunks.

        Args:
            text: Text to split

        Returns:
            List of text chunks
        """
        if not text:
            return []

        # Split by separator first
        splits = text.split(self.separator)

        chunks = []
        current_chunk = []
        current_size = 0

        for split in splits:
            split_size = len(split)

            # If single split is larger than chunk_size, split it further
            if split_size > self.chunk_size:
                # Add current chunk if exists
                if current_chunk:
                    chunks.append(self.separator.join(current_chunk))
                    current_chunk = []
                    current_size = 0

                # Split large text by character
                for i in range(0, split_size, self.chunk_size - self.chunk_overlap):
                    chunks.append(split[i:i + self.chunk_size])

            elif current_size + split_size > self.chunk_size:
                # Current chunk is full, start new chunk
                chunks.append(self.separator.join(current_chunk))

                # Keep overlap
                overlap_text = self.separator.join(current_chunk)
                if len(overlap_text) > self.chunk_overlap:
                    overlap_text = overlap_text[-self.chunk_overlap:]

                current_chunk = [overlap_text, split] if overlap_text else [split]
                current_size = len(self.separator.join(current_chunk))
            else:
                # Add to current chunk
                current_chunk.append(split)
                current_size += split_size + len(self.separator)

        # Add remaining chunk
        if current_chunk:
            chunks.append(self.separator.join(current_chunk))

        return [chunk.strip() for chunk in chunks if chunk.strip()]

    def split_documents(self, documents: List[Document]) -> List[Document]:
        """
        Split multiple documents into chunks.

        Args:
            documents: List of documents to split

        Returns:
            List of chunked documents
        """
        chunked_docs = []

        for doc in documents:
            chunks = self.split_text(doc.content)

            for i, chunk in enumerate(chunks):
                metadata = doc.metadata.copy()
                metadata['chunk_index'] = i
                metadata['total_chunks'] = len(chunks)

                chunked_docs.append(
                    Document(
                        content=chunk,
                        metadata=metadata,
                        source=doc.source
                    )
                )

        return chunked_docs

